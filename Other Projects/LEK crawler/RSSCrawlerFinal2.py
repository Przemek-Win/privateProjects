#!/bin/bash


import pandas as pd
import boto3, botocore, requests, hashlib, pymongo, time, os, datetime, sys, signal, logging, feedparser, json, glob
from selenium import webdriver
from selenium.webdriver.chrome.options import Options  
from selenium.webdriver.common.keys import Keys
from pymongo import MongoClient
from random import randint
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.support import expected_conditions as EC

from docx import Document



global _DRIVER



#Setting up browser for crawler
chrome_options = Options()
chrome_options.add_argument("--headless")
chrome_options.add_argument('--no-sandbox')


chromedriver = os.path.abspath(os.path.dirname(os.path.realpath(__file__))) + '/chromedriver'
os.environ["webdriver.chrome.driver"] = chromedriver
print(os.environ["webdriver.chrome.driver"])

print("Opening browser...") 
_DRIVER = webdriver.Chrome(chromedriver, chrome_options=chrome_options ) #chrome_options=chrome_options




print('Waiting for Scheduler with keyvariable')



def login():
    logging.basicConfig(filename='SmallLogs.log',level=logging.INFO)

    _DRIVER.set_page_load_timeout(120)
    _DRIVER.get('http://egzaminlek.pl/logowanie/')
    #print (_DRIVER.page_source.encode("utf-8"))
    print(_DRIVER.find_element_by_id("username"))

    username=_DRIVER.find_element_by_id("username")
    username.send_keys("XXX")
    password=_DRIVER.find_element_by_id("password")
    password.send_keys("XXX")
    print("Checkpoint")
    _DRIVER.find_element_by_xpath("//button[@type='submit']").click()

    print(_DRIVER.current_url)

    _DRIVER.get('http://egzaminlek.pl/pytania-tematycznie/')
    checkboxes= _DRIVER.find_elements_by_xpath("//input[@type='checkbox'and @name='tags[]' ]")
    tabList = ["#pediatria","#ginekologiaiPoloznictwo","#psychiatria","#chirurgia","#medycynaRodzinna","#intensywnaTerapia"]
    pureList = ["Pediatria","Ginekologia","Psychiatria","Chirurgia","Medycyna rodzinna","Intensywna terapia"]
    
    tabs= _DRIVER.find_elements_by_xpath("//ul[@class='nav nav-tabs']")
    print("Dlugosc", len(checkboxes))
    print("Len", len(tabs))
    _DRIVER.save_screenshot('beforeclicking.png')

    for element in pureList:
        try:
            #@data-toggle=\"tab\" and
            _DRIVER.find_element_by_link_text(element).click()
            #_DRIVER.find_element_by_xpath("//a[ @href=\""+element+"\"]").click()
            _DRIVER.save_screenshot(element+"_Selection.png")
        except Exception as e :
            print("Error with clicking tab", e)
            logging.info('Error occurreed with clicking tab '+ str(e) )
            continue
        for numer in range(0, len(checkboxes)):
            try:
                checkboxes[numer].click()
            except Exception:
                continue
    
    _DRIVER.find_element_by_xpath("//button[@id='wygeneruj-test-button']").click()
    _DRIVER.save_screenshot('afterclicking.png')
    time.sleep(5)
    return _DRIVER



def download_questions():
    #Clicking submit button
    logging.basicConfig(filename='FullLogs.log',level=logging.INFO)
    document = Document()

    _DRIVER=login()
    print("Finished generating test")
    mainDB=pd.DataFrame(columns=['Question and answers', 'Omowienie'])
    print("Beginning the loop")
    for iterator in range(0,4000): #3000
        try:
            time.sleep(2)
            _DRIVER.find_element_by_xpath("//button[@id='pytaniaDropdown']").click()
            #_DRIVER.save_screenshot(str(iterator)+"_selection.png")
            time.sleep(2)
            print("After first screenshot" )
            _DRIVER.find_element_by_xpath(("//li[@data-index='"+ str(iterator+1) + "']")).click()
            #chooser=_DRIVER.find_element_by_xpath(("//li[@data-index='"+ str(iterator+1) + "']"))
            #_DRIVER.execute_script("arguments[0].scrollIntoView();", chooser)
            
            _DRIVER.implicitly_wait(2)
            _DRIVER.execute_script("window.scrollTo(0, 1080)") 
            #_DRIVER.save_screenshot(str(iterator)+"_taking text from questions.png")

            print("After second screenshot" )
            #question=_DRIVER.find_element_by_xpath("//div[@class=\"pytanie\" and @pytanieid=\""+ str(iterator) + "\"]").text
            questions= _DRIVER.find_elements_by_css_selector('p.question-and-answear')
            print("Question  is", questions[2*iterator].text)

            question=questions[2*iterator].text
            
            time.sleep(2)
            #print("Before omowienie",question)
            _DRIVER.find_element_by_xpath(("//button[@class='btn btn-default zobaczOmowienie']")).click()
            #_DRIVER.save_screenshot('screeniefirstafter.png')
            print("After clicking")
            time.sleep(4)
            omowienie=_DRIVER.find_element_by_xpath(("//p[@class='omowienie-body']")).text
            print("After assigning text")
            _DRIVER.save_screenshot('screenieafter.png')
            print("len", str(len(_DRIVER.find_elements_by_xpath(("//button[@class='close']")))))
            _DRIVER.find_elements_by_xpath(("//button[@class='close']"))[1].click()
            time.sleep(2)
            #print("Question is ",  question)
            print("Omowienie is ", type(omowienie),  omowienie)
            print("Iterator at the end",iterator)
            mainDB=mainDB.append({'Question and answers': str(question),'Omowienie': str(omowienie)},ignore_index=True )
            paragraph=document.add_heading("", 0)
            paragraph.add_run("Question" + str(iterator)).bold = True
            
            paragraph=document.add_paragraph("")
            paragraph.add_run("Pytanie i odpowiedzi") .bold = True
            paragraph.add_run(str(question)) 

            paragraph=document.add_paragraph("")
            paragraph.add_run("Omowienie") .bold = True
            paragraph.add_run(str(omowienie))

        except Exception as e:
            print("Error ", e)
            logging.info(str('Error occurreed with question ' + str(iterator) + " Error by the name " +  str(e) ))
            continue
        if (iterator%50 == 0):
            document.save('Q&A.docx')
            mainDB.to_csv(os.path.abspath(os.path.dirname(os.path.realpath(__file__)))+'/WszystkiePytania.csv')
            document = Document('Q&A.docx')
        else:
            continue
    print(mainDB)

    mainDB.to_csv(os.path.abspath(os.path.dirname(os.path.realpath(__file__)))+'/WszystkiePytania.csv')

download_questions()



